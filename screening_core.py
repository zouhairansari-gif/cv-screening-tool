"""
Parsing and Claude API logic shared by the Streamlit app.
Kept separate from the UI code so it can be unit-tested on its own.
"""
import json
import os
import tempfile

import anthropic

MODEL = "claude-sonnet-4-6"


# ---------------------------------------------------------------------------
# File parsing — same approach validated earlier: native extraction first,
# OCR fallback for scanned content, clear warnings rather than silent failure.
# ---------------------------------------------------------------------------

def extract_text_from_uploaded_file(uploaded_file):
    """
    Takes a Streamlit UploadedFile object, writes it to a temp path, and
    extracts plain text. Returns {text, used_ocr, warning}.
    """
    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        tmp_path = tmp.name
    try:
        return extract_text_from_file(tmp_path)
    finally:
        os.unlink(tmp_path)


def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return {"text": f.read(), "used_ocr": False, "warning": None}

    elif ext == ".docx":
        import docx
        try:
            doc = docx.Document(filepath)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            text = "\n".join(p for p in parts if p.strip())
            warning = None if text.strip() else "docx parsed but no text found"
            return {"text": text, "used_ocr": False, "warning": warning}
        except Exception as e:
            return {"text": "", "used_ocr": False, "warning": f"Failed to parse .docx: {e}"}

    elif ext == ".doc":
        return {"text": "", "used_ocr": False,
                "warning": "Legacy .doc isn't supported — please re-save as .docx and re-upload"}

    elif ext == ".pdf":
        import pdfplumber
        text_parts, needs_ocr_pages = [], []
        try:
            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        needs_ocr_pages.append(i)
        except Exception as e:
            return {"text": "", "used_ocr": False, "warning": f"Failed to open PDF: {e}"}

        used_ocr = False
        if needs_ocr_pages:
            try:
                from pdf2image import convert_from_path
                import pytesseract
                images = convert_from_path(filepath)
                ocr_chunks = [pytesseract.image_to_string(images[i]) for i in needs_ocr_pages if i < len(images)]
                if ocr_chunks:
                    text_parts.append("\n".join(ocr_chunks))
                    used_ocr = True
            except Exception:
                pass

        full_text = "\n".join(text_parts)
        warning = None
        if needs_ocr_pages and not used_ocr:
            warning = "Some pages had no extractable text and OCR fallback failed"
        elif not full_text.strip():
            warning = "No text could be extracted from this PDF"
        return {"text": full_text, "used_ocr": used_ocr, "warning": warning}

    elif ext in (".png", ".jpg", ".jpeg"):
        from PIL import Image
        import pytesseract
        try:
            text = pytesseract.image_to_string(Image.open(filepath))
            warning = None if text.strip() else "OCR produced no text — check image quality"
            return {"text": text, "used_ocr": True, "warning": warning}
        except Exception as e:
            return {"text": "", "used_ocr": True, "warning": f"OCR failed: {e}"}

    else:
        return {"text": "", "used_ocr": False, "warning": f"Unsupported file type: {ext}"}


# ---------------------------------------------------------------------------
# Claude API calls
# ---------------------------------------------------------------------------

def get_client(api_key):
    return anthropic.Anthropic(api_key=api_key)


def _extract_json(raw_text):
    """
    Robustly pulls a JSON object out of a Claude response, even if the model
    added conversational text before/after it or wrapped it in a code fence
    that isn't at the very start of the string. Tries, in order:
    1. Parse the raw text directly (the common case).
    2. Find a ```json ... ``` or ``` ... ``` fence anywhere in the text.
    3. Find the first '{' and its matching closing '}' via balanced-brace
       scanning (handles nested objects correctly, unlike a naive rfind).
    Raises a clear, specific error if all three fail — e.g. because the
    response was cut off before finishing (hit the token limit).
    """
    raw_text = raw_text.strip()

    # Attempt 1: parse as-is.
    try:
        return json.loads(raw_text)
    except json.JSONDecodeError:
        pass

    # Attempt 2: extract a fenced code block, wherever it appears.
    import re
    fence_match = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw_text)
    if fence_match:
        try:
            return json.loads(fence_match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Attempt 3: balanced-brace scan from the first '{' to its matching '}'.
    start = raw_text.find("{")
    if start != -1:
        depth = 0
        for i in range(start, len(raw_text)):
            if raw_text[i] == "{":
                depth += 1
            elif raw_text[i] == "}":
                depth -= 1
                if depth == 0:
                    candidate = raw_text[start:i + 1]
                    try:
                        return json.loads(candidate)
                    except json.JSONDecodeError:
                        break

    raise ValueError(
        "Could not parse a valid JSON response from Claude. This usually means the "
        "response was cut off before finishing (try increasing max_tokens) or the "
        "model returned something unexpected. Raw response started with:\n"
        f"{raw_text[:300]}"
    )


def extract_criteria_from_jd(client, jd_text):
    prompt = f"""You are an experienced HR / talent acquisition partner. Given the job
description below, produce a weighted screening rubric with 5-6 criteria that matter
most for this role.

Job description:
{jd_text}

Return ONLY valid JSON in this exact shape, with weights summing to 100:
{{
  "role_title": "short role title inferred from the JD",
  "criteria": [
    {{"name": "...", "weight": 25, "description": "what strong evidence for this looks like"}}
  ]
}}"""
    response = client.messages.create(model=MODEL, max_tokens=1500, messages=[{"role": "user", "content": prompt}])
    return _extract_json(response.content[0].text)


def score_candidate(client, cv_text, criteria, hard_filters=None):
    hard_filters = hard_filters or []

    filters_section = ""
    if hard_filters:
        filters_section = f"""

Also evaluate the candidate against these hard eligibility requirements. These are
pass/fail gates, separate from the scored criteria above — a candidate can score well
overall while still failing one of these.

Hard eligibility requirements:
{json.dumps(hard_filters, indent=2)}

For EACH requirement, return a status of "met", "not_met", or "cannot_determine",
plus a one-sentence rationale citing specific CV evidence (or noting its absence).

IMPORTANT — for any requirement related to work authorization, visa, or nationality:
judge ONLY based on explicit statements in the CV about residency or visa status
(e.g. "holds a valid UAE residence visa," "eligible to work in the UAE without
sponsorship"). NEVER infer this from the candidate's name, the sound of their name,
their country of education, or any other proxy for nationality or ethnicity — if the
CV doesn't explicitly state residency/visa status, mark it "cannot_determine" with the
rationale "not stated in CV — verify directly with the candidate," rather than guessing."""

    prompt = f"""You are scoring a candidate CV against a weighted hiring rubric. Score
based only on evidence in the CV text below — never on assumptions beyond what's written.
If a criterion has no supporting evidence, say so explicitly and score it low rather
than guessing generously.

Rubric criteria:
{json.dumps(criteria, indent=2)}

Candidate CV text:
{cv_text[:12000]}

For EACH criterion, return a score 1 (no evidence) to 5 (strong evidence), plus a
one-sentence rationale citing specific CV evidence (or noting its absence).
{filters_section}

Return ONLY valid JSON in this exact shape:
{{
  "candidate_name": "best-guess name extracted from the CV, or null if not found",
  "scores": [
    {{"criterion": "...", "score": 1, "rationale": "..."}}
  ],
  "filter_results": [
    {{"filter": "...", "status": "met", "rationale": "..."}}
  ]
}}
(Return an empty list for filter_results if no hard requirements were provided.)"""
    response = client.messages.create(model=MODEL, max_tokens=1800, messages=[{"role": "user", "content": prompt}])
    return _extract_json(response.content[0].text)


def extract_glossary(client, jd_text):
    """
    Pulls jargon, acronyms, and industry-specific terms out of a JD that a
    generalist recruiter likely wouldn't know, each with a short plain-English
    definition. Capped so it stays a quick reference, not another report.
    """
    prompt = f"""You are helping a generalist HR recruiter (not a subject-matter expert
in this field) understand a job description before they screen candidates against it.

Job description:
{jd_text}

Identify up to 8 technical terms, acronyms, or industry jargon in this JD that a
generalist recruiter likely would NOT already know (e.g. "RTM", "cost-to-serve",
"EBITDA"). Skip common business terms any recruiter already knows (e.g. "manager",
"stakeholder", "team", "budget").

For each term, give a short, plain-English definition — maximum 20 words — and the
definition itself must not use further jargon.

Return ONLY valid JSON in this exact shape:
{{
  "glossary": [
    {{"term": "...", "definition": "..."}}
  ]
}}
(Return an empty list if the JD has no notable jargon.)"""
    response = client.messages.create(model=MODEL, max_tokens=800, messages=[{"role": "user", "content": prompt}])
    result = _extract_json(response.content[0].text)
    return result.get("glossary", [])[:8]  # defensive cap regardless of model compliance


def rescore_with_comments(client, candidate, criteria, hard_filters, comments):
    """
    Re-scores a candidate using the original CV plus recruiter comments (e.g.
    interview notes, reference-check findings, corrections) as additional
    evidence. The rationale explicitly flags when a comment changed the
    assessment, so a score never shifts without a visible trail explaining why.
    """
    comments_text = "\n".join(f"- {c['text']}" for c in comments if c.get("text", "").strip())

    filters_block = ""
    if hard_filters:
        filters_block = f"""

Also re-evaluate these hard eligibility requirements the same way — weighing both
the CV and the recruiter notes:
{json.dumps(hard_filters, indent=2)}
For EACH, return a status of "met", "not_met", or "cannot_determine", with a
rationale. Apply the same rule as before for work authorization, visa, or
nationality — judge ONLY from explicit statements in the CV or recruiter notes,
NEVER inferred from a name or any other proxy."""

    prompt = f"""You are re-scoring a candidate against a weighted hiring rubric, now
that the recruiter has added notes from interviews, reference checks, or other
verification — not just the original CV.

Rubric criteria:
{json.dumps(criteria, indent=2)}

Original CV text:
{candidate['cv_text'][:12000]}

Recruiter notes (interview findings, verification, corrections):
{comments_text}

Weigh the recruiter notes as authoritative additional evidence — they can confirm,
contradict, or add to what the CV alone shows. For EACH criterion, return an updated
score (1-5) and a rationale. If a recruiter note changed your assessment from what
the CV alone would suggest, say so explicitly (e.g. "Revised down: recruiter
interview note indicates candidate could not provide a specific example").
{filters_block}

Return ONLY valid JSON in this exact shape:
{{
  "scores": [
    {{"criterion": "...", "score": 1, "rationale": "..."}}
  ],
  "filter_results": [
    {{"filter": "...", "status": "met", "rationale": "..."}}
  ]
}}
(Return an empty list for filter_results if no hard requirements were provided.)"""
    response = client.messages.create(model=MODEL, max_tokens=1800, messages=[{"role": "user", "content": prompt}])
    return _extract_json(response.content[0].text)


def merge_requirements_into_rubric(client, role_title, base_criteria, requirements):
    """
    Takes the JD-derived rubric (base_criteria) plus a recruiter's explicit
    requirements — each marked by the recruiter as either a hard filter
    (pass/fail eligibility gate) or a weighted preference (added to the score) —
    and produces a final merged rubric.

    `requirements` is a list of dicts like:
        {"category": "Regional experience", "requirement": "GCC region", "hard_requirement": False}
        {"category": "Work authorization", "requirement": "Valid UAE work visa or transferable", "hard_requirement": True}

    Returns: {"criteria": [...weights summing to 100...], "hard_filters": [...]}
    """
    hard_items = [r for r in requirements if r.get("hard_requirement") and r.get("requirement", "").strip()]
    soft_items = [r for r in requirements if not r.get("hard_requirement") and r.get("requirement", "").strip()]

    # Hard filters don't need an LLM call to define — they're just recorded as-is,
    # and evaluated per-candidate later in score_candidate(). Keeping this
    # deterministic (no LLM judgment on what counts as a "filter") avoids the
    # model quietly reinterpreting a strict requirement as a soft preference.
    hard_filters = [
        {"name": r["category"], "requirement": r["requirement"]}
        for r in hard_items
    ]

    if not soft_items:
        # Nothing to merge into the weighted rubric — return the base criteria untouched.
        return {"criteria": base_criteria, "hard_filters": hard_filters}

    prompt = f"""You are refining a weighted hiring rubric for the role of {role_title}.

Here is the current rubric, derived from the job description:
{json.dumps(base_criteria, indent=2)}

The recruiter has added these additional preferences to fold in (these are things
that should raise or lower a candidate's score, not disqualify them outright):
{json.dumps(soft_items, indent=2)}

Merge these into the rubric:
- If a new preference clearly overlaps with an existing criterion, adjust that
  criterion's description rather than creating a near-duplicate.
- If it's genuinely new, add it as a new criterion.
- Re-weight ALL criteria (existing and new) so they sum to exactly 100, reflecting
  reasonable relative importance — don't just default new items to a small weight.

Return ONLY valid JSON in this exact shape:
{{
  "criteria": [
    {{"name": "...", "weight": 25, "description": "..."}}
  ]
}}"""
    response = client.messages.create(model=MODEL, max_tokens=1500, messages=[{"role": "user", "content": prompt}])
    merged = _extract_json(response.content[0].text)
    return {"criteria": merged["criteria"], "hard_filters": hard_filters}


def generate_interview_questions(client, candidate, role_title):
    """
    Generates probing interview questions targeted at this candidate's most
    consequential weak or uncertain areas — capped at 4 total, kept short
    enough to actually use in an interview, not a report to read beforehand.
    """
    sorted_scores = sorted(candidate["scores"], key=lambda s: s["score"])
    weak_points = sorted_scores[:3]  # the 3 lowest-scoring criteria — most consequential gaps first
    filter_flags = [f for f in candidate.get("filter_results", []) if f["status"] != "met"]

    if all(s["score"] >= 4 for s in candidate["scores"]) and not filter_flags:
        # Strong candidate with nothing to probe defensively — verify depth on
        # the top criteria instead of taking a high score on faith.
        weak_points = sorted(candidate["scores"], key=lambda s: -s["score"])[:2]

    topics = (weak_points + filter_flags)[:4]  # cap input topics so output naturally stays short

    prompt = f"""You are helping a hiring manager prepare for an interview with a
candidate for the role of {role_title}.

Below are the specific areas of this candidate's profile to probe — either scoring
gaps or eligibility items that couldn't be confirmed from the CV alone. Generate ONE
targeted interview question per area, designed to get concrete evidence in the room,
not just re-ask what the CV already says.

Areas to probe:
{json.dumps(topics, indent=2)}

BE EXTREMELY CONCISE — this is a quick reference the interviewer glances at, not a
report to read beforehand. Follow these limits strictly:
- "question": ONE sentence, maximum 20 words.
- "strong_answer_signal": short phrase, maximum 10 words.
- "weak_answer_signal": short phrase, maximum 10 words.

Return ONLY valid JSON in this exact shape:
{{
  "questions": [
    {{
      "topic": "which area this targets",
      "question": "...",
      "strong_answer_signal": "...",
      "weak_answer_signal": "..."
    }}
  ]
}}"""
    response = client.messages.create(model=MODEL, max_tokens=1200, messages=[{"role": "user", "content": prompt}])
    questions = _extract_json(response.content[0].text)["questions"]
    return questions[:4]  # defensive cap regardless of model compliance


def answer_question(client, question, candidates_bundle, role_title, history=None):
    system_prompt = f"""You are helping a hiring manager or recruiter review a candidate
shortlist for the role of {role_title}.

Answer using ONLY the candidate data provided below.
Rules:
- Always name the specific candidate(s) that support your answer.
- If the data doesn't contain enough information, say so explicitly rather than guessing.
- Never invent details not present in the CV text or scores.
- Keep answers concise unless a list is clearly needed.

Candidate data:
{json.dumps(candidates_bundle, indent=2)}"""

    messages = list(history or [])
    messages.append({"role": "user", "content": question})
    response = client.messages.create(model=MODEL, max_tokens=1000, system=system_prompt, messages=messages)
    answer = response.content[0].text
    messages.append({"role": "assistant", "content": answer})
    return answer, messages


def generate_golden_profile(client, role_title, jd_text, criteria, hard_filters):
    """
    Generates a structured 'ideal candidate' profile for the role — organized
    into concrete, labeled categories (not a loose paragraph) so a recruiter
    can source and screen against specifics: what education/certs to look for,
    what industry background, which companies and job titles tend to produce
    strong candidates, and what concrete evidence separates strong from weak.
    Every category is capped short — this is a scan-in-seconds reference,
    not a report.
    """
    prompt = f"""You are an experienced HR / talent acquisition partner with deep
knowledge of this industry. Based on the job description and screening rubric below,
build a structured "ideal candidate" profile for this role.

Role: {role_title}

Job description:
{jd_text}

Screening rubric (for context on what already matters):
{json.dumps(criteria, indent=2)}

Hard eligibility requirements:
{json.dumps(hard_filters, indent=2)}

Be concrete and specific — real qualifications, real company names, real job titles
that would actually appear on a strong candidate's CV for this role and industry.
If a category genuinely doesn't apply to this role (e.g. certifications aren't
typically relevant), return an empty list for it rather than inventing filler.

BE EXTREMELY CONCISE. This will be scanned quickly by a recruiter, not read as a
report. Follow these limits strictly:
- "summary": ONE sentence, maximum 20 words.
- "ideal_education": max 2 short bullets (degree level/field).
- "ideal_certifications": max 2 short bullets (leave empty if not relevant to this role).
- "ideal_industry_experience": max 3 short bullets (specific sub-sector, years, scope).
- "targeted_companies": max 5 real, specific company names a strong candidate is
  likely to have worked at, relevant to this role's industry and region.
- "targeted_similar_roles": max 4 short job titles that map well to this role.
- "evidence_of_strong_experience": max 4 short, concrete phrases describing what a
  strong CV actually shows (not vague traits — specific accomplishments/scope).
- "red_flags": max 3 short bullets.
Each bullet: 4-10 words, not a full sentence.

Return ONLY valid JSON in this exact shape:
{{
  "summary": "...",
  "ideal_education": ["...", "..."],
  "ideal_certifications": ["...", "..."],
  "ideal_industry_experience": ["...", "..."],
  "targeted_companies": ["...", "..."],
  "targeted_similar_roles": ["...", "..."],
  "evidence_of_strong_experience": ["...", "..."],
  "red_flags": ["...", "..."]
}}"""
    response = client.messages.create(model=MODEL, max_tokens=1000, messages=[{"role": "user", "content": prompt}])
    profile = _extract_json(response.content[0].text)
    # Defensive caps regardless of whether the model followed the limits —
    # keeps this crisp even if a response comes back longer than asked.
    caps = {
        "ideal_education": 2, "ideal_certifications": 2, "ideal_industry_experience": 3,
        "targeted_companies": 5, "targeted_similar_roles": 4,
        "evidence_of_strong_experience": 4, "red_flags": 3,
    }
    for field, limit in caps.items():
        profile[field] = profile.get(field, [])[:limit]
    return profile


def compare_candidate_to_golden_profile(client, candidate, golden_profile, role_title):
    """
    Compares one candidate's actual CV against the structured golden profile —
    checking specifically against its labeled categories (education, industry
    experience, target companies, target roles, evidence) rather than a vague
    overall impression. Deliberately kept short: a scannable fit-level tag plus
    a few short bullets, not a written comparison.
    """
    prompt = f"""You are comparing a candidate's CV against a structured ideal
("golden") candidate profile for the role of {role_title}.

Golden profile:
{json.dumps(golden_profile, indent=2)}

Candidate CV text:
{candidate['cv_text'][:12000]}

Check the candidate specifically against the golden profile's categories: education,
certifications, industry experience, target companies (did they work at one of the
listed companies or a clear peer?), target similar roles (did they hold one of these
titles or equivalent?), and the concrete evidence markers. Cite CV evidence; never
invent details not present in the CV.

BE EXTREMELY CONCISE. This will be scanned quickly by a recruiter, not read as a
report. Follow these limits strictly:
- "one_liner": ONE short sentence, maximum 15 words.
- "matches": maximum 3 bullets, each a short phrase of 6-10 words, naming which
  category it satisfies where relevant (e.g. "Worked at Unilever, a target company").
- "gaps": maximum 3 bullets, each a short phrase of 6-10 words.

Return ONLY valid JSON in this exact shape:
{{
  "fit_level": "Strong", "Moderate", or "Weak",
  "one_liner": "...",
  "matches": ["...", "..."],
  "gaps": ["...", "..."]
}}"""
    response = client.messages.create(model=MODEL, max_tokens=600, messages=[{"role": "user", "content": prompt}])
    comparison = _extract_json(response.content[0].text)
    comparison["matches"] = comparison.get("matches", [])[:3]
    comparison["gaps"] = comparison.get("gaps", [])[:3]
    return comparison


def normalize_weights(criteria):
    """
    After a recruiter manually edits criterion weights, they may no longer sum
    to exactly 100. This rescales them proportionally so they do, preserving
    the relative importance the recruiter set rather than silently overriding it.
    Rounds each weight to a whole number, then adjusts the last item so the
    total is exactly 100 even after rounding (rounding each independently can
    otherwise leave the total a fraction off, e.g. 99.9 instead of 100).
    """
    total = sum(c["weight"] for c in criteria)
    if total == 0:
        # Avoid division by zero — fall back to equal weighting.
        base = 100 // len(criteria)
        remainder = 100 - base * len(criteria)
        return [
            {**c, "weight": base + (1 if i < remainder else 0)}
            for i, c in enumerate(criteria)
        ]

    scaled = [round(c["weight"] * 100 / total) for c in criteria]
    diff = 100 - sum(scaled)
    scaled[-1] += diff  # absorb any rounding remainder into the last item
    return [{**c, "weight": w} for c, w in zip(criteria, scaled)]


def _pdf_safe(text):
    """
    The default PDF font (Helvetica) only supports Latin-1 characters, but
    Claude's output commonly includes em-dashes, curly quotes, and similar
    punctuation that Latin-1 doesn't cover. Replaces the common cases with
    plain-ASCII equivalents, then strips anything else that still doesn't fit
    rather than letting the whole PDF generation crash on one character.
    """
    if not text:
        return ""
    replacements = {
        "\u2014": "-", "\u2013": "-",       # em dash, en dash
        "\u2018": "'", "\u2019": "'",       # curly single quotes
        "\u201c": '"', "\u201d": '"',       # curly double quotes
        "\u2026": "...",                     # ellipsis
        "\u2022": "-",                       # bullet
    }
    for original, replacement in replacements.items():
        text = text.replace(original, replacement)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_candidate_pdf(candidate, role_title, criteria, golden_profile=None, interview_questions=None):
    """
    Builds a one-page(ish) PDF summary for a single candidate: weighted score,
    per-criterion rationale, eligibility flags, golden profile comparison (if
    available), and interview questions (if generated). Returns raw PDF bytes,
    suitable for st.download_button.
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _pdf_safe(candidate["candidate_name"]), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(90, 90, 90)
    pdf.cell(0, 7, _pdf_safe(f"{role_title}  -  Weighted score: {candidate['weighted_score']}/5"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    filter_results = candidate.get("filter_results", [])
    if filter_results:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Eligibility requirements", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for f in filter_results:
            status_label = {"met": "MET", "not_met": "NOT MET", "cannot_determine": "VERIFY DIRECTLY"}.get(
                f["status"], f["status"].upper()
            )
            pdf.multi_cell(0, 6, _pdf_safe(f"[{status_label}] {f['filter']}: {f['rationale']}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Criterion scores", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for s in candidate["scores"]:
        pdf.multi_cell(0, 6, _pdf_safe(f"[{s['score']}/5] {s['criterion']}: {s['rationale']}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    if golden_profile is not None and candidate.get("golden_profile_comparison"):
        comp = candidate["golden_profile_comparison"]
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, _pdf_safe(f"Golden profile fit: {comp['fit_level']}"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, _pdf_safe(comp["one_liner"]), new_x="LMARGIN", new_y="NEXT")
        if comp.get("matches"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, "Matches:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            for m in comp["matches"]:
                pdf.multi_cell(0, 6, _pdf_safe(f"+ {m}"), new_x="LMARGIN", new_y="NEXT")
        if comp.get("gaps"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, "Gaps:", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            for g in comp["gaps"]:
                pdf.multi_cell(0, 6, _pdf_safe(f"- {g}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    if interview_questions:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Suggested interview questions", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for q in interview_questions:
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 6, _pdf_safe(f"Q: {q['question']}"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, _pdf_safe(f"   Strong answer: {q['strong_answer_signal']}"), new_x="LMARGIN", new_y="NEXT")
            pdf.multi_cell(0, 6, _pdf_safe(f"   Weak answer: {q['weak_answer_signal']}"), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

    return bytes(pdf.output())


def weighted_score(scores, criteria):
    total = sum(
        s["score"] * next(c["weight"] for c in criteria if c["name"] == s["criterion"])
        for s in scores
    ) / 100
    return round(total, 2)
